import pandas as pd
from docx import Document
from bs4 import BeautifulSoup
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LCDocument
from pathlib import Path

def load_documents(file_paths):
    """支持 pdf, txt, md, docx, html, csv"""
    docs = []
    for path in file_paths:
        path = Path(path)
        suffix = path.suffix.lower()
        try:
            if suffix == '.pdf':
                loader = PyPDFLoader(str(path))
                docs.extend(loader.load())
            elif suffix in ['.txt', '.md']:
                loader = TextLoader(str(path), encoding='utf-8')
                docs.extend(loader.load())
            elif suffix == '.docx':
                doc = Document(path)
                text = "\n".join([para.text for para in doc.paragraphs])
                docs.append(LCDocument(page_content=text, metadata={"source": str(path)}))
            elif suffix == '.html':
                with open(path, 'r', encoding='utf-8') as f:
                    soup = BeautifulSoup(f, 'html.parser')
                    text = soup.get_text(separator='\n')
                docs.append(LCDocument(page_content=text, metadata={"source": str(path)}))
            elif suffix == '.csv':
                df = pd.read_csv(path)
                text = df.to_string(index=False)
                docs.append(LCDocument(page_content=text, metadata={"source": str(path)}))
            else:
                print(f"警告：不支持的文件格式 {suffix}，已跳过 {path}")
        except Exception as e:
            print(f"读取文件 {path} 出错：{e}")
    return docs

def chunk_documents(docs, chunk_size=500, chunk_overlap=50):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", "。", "！", "？", "；", " ", ""]
    )
    return splitter.split_documents(docs)